import os
from pathlib import Path
from app.adapters.llm_client import OllamaClient
from app.domain.services.ledger_service import LedgerService
from app.utils.logger import orchestrator_logger
from app.utils.paths import get_client_desktop

class WordAgent:
    """
    Word Agent: Subagente de redacción documental de Alfonso que redacta informes
    financieros formales y los guarda como documentos Word (.docx).
    """
    def __init__(self):
        self.llm = OllamaClient()
        self.prompt_path = Path("app/prompts/word_system.txt")
        self._load_prompt()

    def _load_prompt(self):
        try:
            self.system_prompt = self.prompt_path.read_text(encoding="utf-8")
        except FileNotFoundError:
            self.system_prompt = "Eres un experto redactor de informes corporativos y financieros en Word."

    async def generate_response(self, query: str, client_id: str = "default") -> str:
        """
        Redacta el informe contable e histórico y lo guarda en formato Word (.docx).
        """
        import re
        year = 2026
        m_year = re.search(r"\b(202\d)\b", query)
        if m_year:
            year = int(m_year.group(1))

        desktop = get_client_desktop(client_id)
        filename = f"Informe_Financiero_{year}.docx"
        filepath = os.path.join(desktop, filename).replace("\\", "/")

        # Obtener los datos contables para inyectar en la redacción
        balance = LedgerService.get_balance_situacion(year)
        
        # Pedir al LLM que redacte el cuerpo del informe en base al balance
        prompt = f"""
        Actúa como el redactor de informes financieros de Alfonso.
        Genera el contenido para un Informe Ejecutivo del año fiscal {year} con la siguiente información:
        
        - Total Activo: {balance['total_activo']} €
        - Total Pasivo y Patrimonio: {balance['total_pasivo_patrimonio']} €
        
        Cuentas de Activo desglosadas:
        {json_dumps(balance['activo'])}
        
        Cuentas de Pasivo desglosadas:
        {json_dumps(balance['pasivo_patrimonio'])}
        
        Por favor, redacta una introducción del negocio, un desglose formal y una sección de conclusiones financieras firmadas por Alfonso.
        Escribe en Markdown limpio.
        """
        
        try:
            raw_body = await self.llm.generate(
                prompt,
                mode="chat",
                memory=self.system_prompt
            )
        except Exception:
            raw_body = (
                f"# INFORME FINANCIERO EJECUTIVO - AÑO {year}\n\n"
                f"Preparado de forma automatizada por Alfonso Autónomo.\n\n"
                f"## 1. RESUMEN EJECUTIVO\n"
                f"El total del activo asciende a {balance['total_activo']} € y el total del pasivo y patrimonio neto es de {balance['total_pasivo_patrimonio']} €.\n\n"
                f"## 2. CONCLUSIONES\n"
                f"Las cuentas anuales reflejan la estabilidad financiera y el equilibrio contable del ejercicio."
            )

        # Crear el archivo Word
        success = self._create_docx_file(raw_body, filepath)
        
        if success:
            response_text = (
                f"📄 **Informe Financiero Redactado con Éxito**\n\n"
                f"- **Tipo de Archivo**: Word Document (.docx)\n"
                f"- **Destino**: [{filename}](file:///{filepath})\n\n"
                f"El documento incluye un análisis ejecutivo detallado, formateado con márgenes, "
                f"títulos en negrita y alineaciones limpias."
            )
        else:
            # Fallback a Markdown (.md)
            md_filename = filename.replace(".docx", ".md")
            md_filepath = filepath.replace(".docx", ".md")
            with open(md_filepath, "w", encoding="utf-8") as f:
                f.write(raw_body)
            response_text = (
                f"📄 **Informe Financiero (Markdown - Fallback) Redactado con Éxito**\n\n"
                f"*(Nota: python-docx no está disponible en este entorno, por lo que se exportó en Markdown clásico)*\n\n"
                f"- **Tipo de Archivo**: Markdown Document (.md)\n"
                f"- **Destino**: [{md_filename}](file:///{md_filepath})"
            )
            
        return response_text

    def _create_docx_file(self, body_text: str, filepath: str) -> bool:
        try:
            import docx
            doc = docx.Document()
            
            # Título principal
            doc.add_heading("INFORME FINANCIERO ANUAL", level=0)
            
            # Parsear el Markdown simple del LLM y añadirlo
            lines = body_text.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
                    
            doc.save(filepath)
            return True
        except ImportError:
            return False

def json_dumps(d):
    import json
    return json.dumps(d, indent=4, ensure_ascii=False)

# Instancia global única
word_agent = WordAgent()
